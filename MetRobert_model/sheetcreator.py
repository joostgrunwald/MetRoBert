# import gspread
# from oauth2client.service_account import ServiceAccountCredentials
# from pprint import pprint

import os
from docx.enum.text import WD_COLOR_INDEX

# get current location folder
location = os.path.realpath(
    os.path.join(os.getcwd(), os.path.dirname(__file__)))

###############
# WORD OUTPUT #
###############

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('MetRuBert Output' , 0)

######################
# SENTENCE RETRIEVAL #
######################
                     
previous_sentence = None
pos_list = []
met_list = []

sentences = []

with open(os.path.join(location, 'output.tsv'), mode='r', encoding='utf-8') as output:
    for line in output:
        if line[:5] != "index":
            splitted = line.split("\t")

            sentence = splitted[1]

            if sentence != previous_sentence != None:

                #TODO: wrapup
                #print(previous_sentence)
                #print(met_list)
                #print(pos_list)

                split_sen = previous_sentence.split()
                p = document.add_paragraph('')
                     
                for i in range(len(split_sen)):
                     word = split_sen[i]
                     if str(i) + ' ' in met_list:
                         #print("reached")
                         p.add_run(word).bold = True
                     else:
                           for pos in pos_list:
                               #print(pos)
                               if pos[0] == str(i) + ' ':
                                   if pos[1] == "NOUN":
                                       font = p.add_run(word).font
                                       font.highlight_color = WD_COLOR_INDEX.YELLOW
                                   else:
                                       p.add_run(word)
                               else:
                                   p.add_run(word)
                     p.add_run(' ')
                sentences.append(previous_sentence)

                pos_list = []
                met_list = []

            else:
                # append pos tag
                pos = splitted[2]
                index = splitted[3]
                pos_list.append([index, pos])

                # append metaphor
                met = splitted[5]
                if met.find("1") != -1 and met.find("-1") == -1:
                    met_list.append(index)

            previous_sentence = sentence

    

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_page_break()

document.save('MetRubert_runx.docx')

#####################
# CREATE EXCEL FILE #
#####################

# # import xlsxwriter module
# import xlsxwriter
 
# # Workbook() takes one, non-optional, argument
# # which is the filename that we want to create.
# workbook = xlsxwriter.Workbook('MetRuBert_run_x.xlsx')
 
# # The workbook object is then used to add new
# # worksheet via the add_worksheet() method.
# worksheet = workbook.add_worksheet()
 
# #HEADER
# worksheet.write("A1", "index")
# worksheet.write("B1", "sentence")
# worksheet.write("C1", "confidence list")

# for i in range(len(sentences)):
#     indexcell = 'A' + str(i+2)
#     sentencecell = 'B' + str(i+2)
    
#     worksheet.write(indexcell, i+1)
#     worksheet.write(sentencecell, sentences[i])
 
# # Finally, close the Excel file
# # via the close() method.
# workbook.close()

# #######################
# # FURTHER ADJUST FILE #
# #######################
# import openpyxl
 
# # Give the location of the file
# path = os.path.join(location, "MetRuBert_run_x.xlsx")
 
# # To open the workbook
# # workbook object is created
# wb_obj = openpyxl.load_workbook(path)
 
# # Get workbook active sheet object
# # from the active attribute
# sheet_obj = wb_obj.active

# #set title
# sheet_obj.title = "MetRubert_sheet"

# #make column b wider
# sheet_obj.column_dimensions['B'].width = 180

# wb_obj.save("MetRuBert_run_x.xlsx")

#! BELOW IS OUTDATED GOOGLE SHEETS CODE, WE USE EXCEL INSTEAD
# create worksheet
# worksheet = out.add_worksheet(
#     title="Primary worksheet", rows=len(sentences)+1, cols=20)

# # delete old worksheet
# #out.reorder_worksheets([worksheet, sheet1])

# worksheet.append_rows(values=[["index", "sentence", "confidence list"]])
# for i in range(len(sentences)):
#     worksheet.append_rows(values=[[i, sentences[i]]])


# switch = False
# for i in range(len(sentences)):
#     if switch == False:
#         switch = True
#     else:
#         switch = False
#         ran = f"A{str(i)}:C{str(i)}"
#         worksheet.format(ran,
#                          {
#                              "backgroundColor": {
#                                  "red": 0.911,
#                                  "green": 0.909,
#                                  "blue": 0.987,
#                                  "alpha": 1.0
#                              },
#                              "textFormat": {
#                                  "fontFamily": "georgia",
#                              }
#                          })

# # format the top
# worksheet.format("A1:C1",
#                  {
#                      "backgroundColor": {
#                          "red": 0.537,
#                          "green": 0.537,
#                          "blue": 0.922,
#                          "alpha": 1.0
#                      },
#                      "textFormat": {
#                          "fontFamily": "georgia",
#                          "bold": True
#                      }
#                  })
# TODO: open output.tsv file

# TODO: save panda dataframe to google sheets

# TODO: make metpahors bold, use background color for pos tags, rewrite softmax to confidence

# sheet = client.open("MetRubert_Sheets").sheet1  # Open the spreadhseet

# data = sheet.get_all_records()  # Get a list of all records
# print(data)

# row = sheet.row_values(3)  # Get a specific row
# col = sheet.col_values(3)  # Get a specific column
# cell = sheet.cell(1,2).value  # Get the value of a specific cell

# insertRow = ["hello", 5, "red", "blue"]
# sheet.appendRow(insertRow, 4)  # Insert the list as a row at index 4

# sheet.update_cell(2,2, "CHANGED")  # Update one cell

# numRows = sheet.row_count  # Get the number of rows in the sheet
